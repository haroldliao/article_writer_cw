import re
import json
from io import BytesIO
from docx import Document


def sanitize_markdown(md: str) -> str:
    """基本清理：行尾空白、多餘空行、引號、標題前空行等"""
    if not md:
        return ""
    
    # 統一換行
    md = md.replace("\r\n", "\n").replace("\r", "\n")
    
    # 中文引號統一：「」
    md = md.replace("“", "「").replace("”", "」")
    
    # 行尾空白
    md = re.sub(r"[ \t]+$", "", md, flags=re.MULTILINE)
    
    # 連續 3 個以上空行 → 2 個
    md = re.sub(r"\n{3,}", "\n\n", md)
    
    # 確保每個標題（#/##/### ）前有一個空行（檔首除外）
    md = re.sub(r"([^\n])\n(#{1,3} )", r"\1\n\n\2", md)
    
    return md.strip()


def count_words(md: str) -> int:
    """
    統計中英文字數
    - 中文字：每個字計為 1
    - 英文詞：以空白分隔計數
    - 排除 Markdown 標記符號
    """
    if not md:
        return 0
    
    # 移除 Markdown 標記（標題符號、粗體、斜體等）
    text = re.sub(r"^#{1,6}\s+", "", md, flags=re.MULTILINE)  # 標題
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # 粗體
    text = re.sub(r"\*(.+?)\*", r"\1", text)  # 斜體
    text = re.sub(r"`(.+?)`", r"\1", text)  # 行內程式碼
    
    # 統計中文字（CJK 統一表意文字）
    chinese_count = len(re.findall(r"[\u4e00-\u9fff]", text))
    
    # 統計英文詞（連續字母序列）
    english_count = len(re.findall(r"\b[a-zA-Z]+\b", text))
    
    return chinese_count + english_count


def count_paragraphs(md: str) -> int:
    """
    統計段落數（以雙換行符分隔）
    - 排除空段落
    - 排除純標題段落
    """
    if not md:
        return 0
    
    blocks = md.split("\n\n")
    
    paragraphs = [
        block.strip() 
        for block in blocks 
        if block.strip() and not re.match(r"^#{1,6}\s+", block.strip())
    ]
    
    return len(paragraphs)


def count_quotes(md: str) -> int:
    """
    統計引號數量（以成對為一組）
    支援：「...」、"..."、'...'
    """
    if not md:
        return 0
    
    pairs = 0
    pairs += len(re.findall(r"「[^」]+」", md))
    pairs += len(re.findall(r"\"[^\"]+\"", md))
    pairs += len(re.findall(r"'[^']+'", md))
    return pairs


def extract_headings(md: str) -> list[str]:
    """抓取所有小標題（### 開頭）"""
    return [
        line[4:].strip() 
        for line in md.splitlines() 
        if line.startswith("### ")
    ]


def extract_all_headings(md: str) -> dict[str, list[str]]:
    """
    提取所有層級的標題
    返回：{"h1": [...], "h2": [...], "h3": [...]}
    """
    headings = {"h1": [], "h2": [], "h3": []}
    
    for line in md.splitlines():
        if line.startswith("### "):
            headings["h3"].append(line[4:].strip())
        elif line.startswith("## "):
            headings["h2"].append(line[3:].strip())
        elif line.startswith("# "):
            headings["h1"].append(line[2:].strip())
    
    return headings


def analyze_article(md: str, word_range: tuple[int, int] = (1500, 2000), 
                   min_quotes: int = 5) -> dict:
    """
    完整分析 Markdown 文章
    """
    word_count = count_words(md)
    paragraphs = count_paragraphs(md)
    quotes = count_quotes(md)
    headings = extract_all_headings(md)
    
    return {
        "word_count": word_count,
        "within_range": word_range[0] <= word_count <= word_range[1],
        "word_range": word_range,
        "paragraphs": paragraphs,
        "quotes": quotes,
        "has_enough_quotes": quotes >= min_quotes,
        "min_quotes": min_quotes,
        "headings": headings,
        "h3_count": len(headings["h3"]),
        "total_heading_count": sum(len(v) for v in headings.values())
    }


def basic_summary_text(checks: dict, retries: int = 0) -> str:
    """將檢查結果轉成簡要文字摘要"""
    word_range = checks.get('word_range', (1500, 2000))
    
    lines = [
        f"字數：{checks.get('word_count', 0)}（{word_range[0]}–{word_range[1]} 合格：{checks.get('within_range', False)}）",
        f"段落數（以空行計）：{checks.get('paragraphs', 0)}",
        f"引號數（成對計算）：{checks.get('quotes', 0)}（≥{checks.get('min_quotes', 5)} 合格：{checks.get('has_enough_quotes', False)}）",
        f"小標題數（###）：{checks.get('h3_count', 0)}",
        f"總標題數：{checks.get('total_heading_count', 0)}",
        f"自動修稿次數：{retries}",
    ]
    
    return "\n".join(lines)


def build_docx_from_markdown(md: str) -> bytes:
    """
    輕量級 Markdown -> DOCX
    支援：標題(#/##/###)、段落、空行
    """
    doc = Document()
    
    for raw in md.splitlines():
        line = raw.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        else:
            doc.add_paragraph(line)
    
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_meta_json(
    subject: str,
    company: str,
    people: str,
    participants: str,
    article_md: str,
    checks: dict,
    retries: int,
    word_count_range: tuple[int, int] | None = None,
    paragraphs: int | None = None
) -> bytes:
    """輸出一份結構化的 meta.json（UTF-8 bytes）"""
    meta = {
        "subject": subject,
        "company": company,
        "people": people,
        "participants": participants,
        "headings": checks.get("headings", {}),
        "checks": checks,
        "auto_edit_retries": retries,
        "constraints": {
            "word_count_range": word_count_range,
            "min_paragraphs": paragraphs
        },
        # 用 word_count 取代字元長度
        "article_word_count": checks.get("word_count", 0)
    }
    
    return json.dumps(meta, ensure_ascii=False, indent=2).encode("utf-8")


# 使用範例
if __name__ == "__main__":
    sample_md = """# 主標題

## 副標題

### 小標題一

這是第一個段落,包含「中文引號」和一些 English words。
這是同一段落的延續。

### 小標題二

這是第二個段落,又有「另一組引號」。

這是第三個段落,包含 "English quotes" 測試。
"""
    
    cleaned = sanitize_markdown(sample_md)
    analysis = analyze_article(cleaned, word_range=(50, 200), min_quotes=2)
    print(basic_summary_text(analysis, retries=0))
    print("\n詳細分析：")
    print(json.dumps(analysis, ensure_ascii=False, indent=2))
